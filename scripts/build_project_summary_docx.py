from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "Dokument_Permbledhes_Projektit_EnerCo.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=11, color="000000", bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Faqe ")
    set_run_font(run, size=9, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=9.5, color=NAVY)


def add_step(doc, number, title, objective, implemented, outputs, results, challenge):
    doc.add_heading(f"{number}. {title}", level=1)
    add_body(doc, f"Qëllimi. {objective}", "Qëllimi.")
    doc.add_heading("Çfarë u implementua", level=2)
    add_bullets(doc, implemented)
    doc.add_heading("Output-et kryesore", level=2)
    for output in outputs:
        add_code(doc, output)
    if results:
        doc.add_heading("Rezultatet kryesore", level=2)
        add_bullets(doc, results)
    doc.add_heading("Sfida dhe zgjidhja", level=2)
    add_body(doc, challenge)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "ENERCO  |  PËRMBLEDHJE E PROJEKTIT ANALITIK"
    set_run_font(header.runs[0], size=9, color=MID_GRAY, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    # Editorial cover.
    doc.add_paragraph().paragraph_format.space_after = Pt(70)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("DOKUMENT PËRMBLEDHËS")
    set_run_font(run, size=11, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(16)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Analiza e profileve të konsumit të energjisë elektrike")
    set_run_font(run, size=28, color=NAVY, bold=True)
    title.paragraph_format.space_after = Pt(12)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Procesi i zhvillimit, rezultatet dhe sfidat kryesore")
    set_run_font(run, size=15, color=BLUE)
    subtitle.paragraph_format.space_after = Pt(38)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("EnerCo  |  Qershor 2025 - Qershor 2026")
    set_run_font(run, size=11, color=MID_GRAY, bold=True)
    meta.paragraph_format.space_after = Pt(8)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Version i anonimizuar - kompanitë paraqiten si “Kompania N”")
    set_run_font(run, size=10, color=MID_GRAY, italic=True)

    doc.add_page_break()
    doc.add_heading("Përmbledhje ekzekutive", level=1)
    add_body(
        doc,
        "Projekti u ndërtua si një pipeline Python i përsëritshëm për kontrollin, "
        "përpunimin dhe interpretimin e të dhënave orare të konsumit të energjisë. "
        "Ai e shndërron workbook-un burimor në raporte të kontrolluara, profile për "
        "kompani dhe njehsor, outlierë të interpretuar, analizë moti/festash, grupe "
        "sjelljeje dhe një ndërfaqe interaktive Streamlit.",
    )
    add_body(
        doc,
        "Parimi qendror. Të dhënat dhe rezultatet mbeten anonime; konsumi A+ dhe "
        "injektimi A- ruhen veçmas; asnjë outlier nuk fshihet automatikisht; çdo hap "
        "është i testuar dhe i dokumentuar.",
        "Parimi qendror.",
    )

    doc.add_heading("Statusi i fazave", level=2)
    rows = [
        ("Faza 0", "Struktura teknike dhe konfigurimi", "Përfunduar"),
        ("Hapi 1", "Kontrolli i cilësisë", "Përfunduar"),
        ("Hapi 2", "Transformimi në format të gjatë", "Përfunduar"),
        ("Hapi 3", "Metrikat e profilit", "Përfunduar"),
        ("Hapi 4", "Outlierët", "Përfunduar; sektori në pritje"),
        ("Hapi 5", "Moti dhe festat", "Përfunduar"),
        ("Hapi 6", "Grupimi i kompanive", "Përfunduar"),
        ("UI", "Ndërfaqja Streamlit", "Përfunduar"),
        ("Run-all", "Orkestrimi i sigurt", "Përfunduar"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Faza", "Përshkrimi", "Statusi"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, BLUE)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, size=10, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for phase, description, status in rows:
        cells = table.add_row().cells
        for i, value in enumerate((phase, description, status)):
            run = cells[i].paragraphs[0].add_run(value)
            set_run_font(run, size=9.5)
        if len(table.rows) % 2 == 0:
            for cell in cells:
                shade_cell(cell, LIGHT_GRAY)
    set_table_geometry(table, [1500, 4860, 3000])

    add_step(
        doc, "1", "Faza 0 - Struktura teknike", 
        "Krijimi i një baze të qartë, të konfigurueshme dhe të ripërdorshme për gjithë analizën.",
        [
            "U krijuan dosjet config, data/input, data/external, data/interim, data/processed, outputs, src dhe tests.",
            "U shtuan pyproject.toml, konfigurimi project.json, CLI-ja dhe politika e privatësisë.",
            "Inputet konfidenciale dhe output-et u përjashtuan nga versionimi publik.",
        ],
        ["config/project.json", "src/enerco_analysis/", "tests/"],
        [],
        "Sfida ishte ndarja e qartë mes inputit konfidencial, rezultateve teknike, raporteve për dorëzim dhe kodit. Struktura me shtresa e zgjidhi këtë dhe mundësoi auditim të çdo faze.",
    )

    add_step(
        doc, "2", "Hapi 1 - Kontrolli i cilësisë",
        "Verifikimi i integritetit të leximeve para çdo llogaritjeje analitike.",
        [
            "Kontrolle për mungesa, negative, vlera jonumerike, zero të vazhdueshme, skaje ekstreme dhe boshtin kohor.",
            "Dy statuse cilësie: për gjithë periudhën dhe për periudhën aktive të njehsorit.",
            "Pragjet kryesore: mbi 10% mungesa, zero për të paktën 48 orë dhe vlerë mbi 50 herë mesataren.",
        ],
        ["outputs/data_quality_report.xlsx", "data/interim/meter_quality.parquet"],
        [
            "6 sheet-e dhe 9,480 orë për sheet.",
            "81 kompani, 499 njehsorë fizikë dhe 517 seri energjie.",
            "0 negative, 0 jonumerike, 369 sekuenca zero >=48 orë dhe 11 skaje ekstreme.",
        ],
        "Shumë njehsorë kishin fillim të vonë dhe dilnin gabimisht si të papërdorshëm. Statusi i periudhës aktive ndau mungesat para aktivizimit nga boshllëqet reale gjatë operimit.",
    )

    add_step(
        doc, "3", "Hapi 2 - Transformimi në format të gjatë",
        "Shndërrimi i workbook-ut të gjerë në një strukturë analitike me një rresht për seri dhe orë.",
        [
            "U ruajtën identiteti anonim, njehsori, data, ora, tarifa, drejtimi i energjisë dhe kWh.",
            "A+ u klasifikua si consumption_import dhe A- si injection_export.",
            "Konsumi dhe injektimi nuk netohen.",
        ],
        ["data/processed/hourly_consumption_long.parquet"],
        ["Rreth 4.9 milionë rreshta të strukturuar për analizë efikase."],
        "Sheet-i Prosumer kishte dy kolona për të njëjtin njehsor. Zgjidhja ishte ruajtja e të njëjtit meter_id me dy energy_flow të ndryshme, pa shenja negative dhe pa humbur injektimin.",
    )

    add_step(
        doc, "4", "Hapi 3 - Metrikat e profilit",
        "Përshkrimi i formës reale të konsumit për çdo njehsor dhe kompani.",
        [
            "U llogaritën raporti pik/jo-pik, raporti javë/fundjavë, CV dhe load factor.",
            "U llogaritën indekset verë/dimër, etiketa e sezonalitetit dhe trendi mujor.",
            "U shtuan flamuj të besueshmërisë sipas mbulimit 90% dhe analiza e heterogjenitetit të njehsorëve.",
        ],
        ["outputs/profile_metrics.xlsx"],
        ["499 njehsorë fizikë, 517 seri dhe 81 profile kompanish."],
        "Totali i kompanisë mund të fshehë profile të kundërta brenda saj. Korrelacionet e profileve 24-orëshe dhe diapazonet e metrikave zbulojnë njehsorët pjesërisht ose dukshëm të ndryshëm.",
    )

    add_step(
        doc, "5", "Hapi 4 - Identifikimi i outlierëve",
        "Gjetja e orëve dhe profileve që devijojnë ndjeshëm nga sjellja historike.",
        [
            "Z-score llogaritet si (vlera - mesatarja) / devijimi standard.",
            "Orët me vlerë absolute Z mbi 3 sinjalizohen, por nuk fshihen.",
            "Rekomandimi ndahet mes shqyrtimit teknik dhe verifikimit të sjelljes së biznesit.",
        ],
        ["outputs/outlier_report_enriched.xlsx"],
        [
            "45,632 outlierë në nivel serie/njehsori.",
            "5,840 outlierë në nivel kompanie.",
            "11 raste kalojnë edhe pragun teknik 50 herë mesataren.",
        ],
        "Outlier nuk do të thotë domosdoshmërisht defekt. Mund të jetë aktivitet real, ndërprerje, festë ose temperaturë ekstreme. Krahasimi zyrtar brenda sektorit mbetet i bllokuar sepse mungon metadata e veprimtarisë.",
    )

    add_step(
        doc, "6", "Hapi 5 - Moti dhe festat",
        "Shtimi i kontekstit të jashtëm për interpretim më të saktë të konsumit dhe outlierëve.",
        [
            "U përdorën temperaturat historike të Prishtinës dhe festat zyrtare të Kosovës.",
            "HDD18 = max(18°C - temperatura mesatare, 0).",
            "CDD18 = max(temperatura mesatare - 18°C, 0).",
            "Efekti i festës krahasohet me ditë normale të së njëjtës ditë jave.",
        ],
        ["outputs/weather_holiday_analysis.xlsx", "outputs/outlier_report_enriched.xlsx"],
        [
            "81 profile kompanish me analizë moti dhe 1,259 raste kompani-festë.",
            "Outlierëve iu shtua konteksti: festë, temperaturë e ulët, e lartë ose pa kontekst të drejtpërdrejtë.",
        ],
        "Lokacioni individual i kompanive mungon, ndaj Prishtina përdoret si proxy. Korrelacioni nuk interpretohet si shkakësi dhe festa zyrtare nuk nënkupton automatikisht mbyllje të biznesit.",
    )

    add_step(
        doc, "7", "Hapi 6 - Grupimi i kompanive",
        "Zbulimi i grupeve natyrore të konsumit sipas sjelljes reale dhe jo vetëm etiketës së biznesit.",
        [
            "U përdorën gjashtë metrika: pik/jo-pik, javë/fundjavë, CV, load factor, indeks veror dhe dimëror.",
            "U testuan k=2 deri 10, Elbow, Silhouette, PCA dhe tre metoda shkallëzimi.",
            "MinMaxScaler u zgjodh për të kufizuar ndikimin e vlerave ekstreme.",
        ],
        ["outputs/company_clustering.xlsx", "outputs/charts/clustering_pca.png"],
        [
            "79 kompani të grupuara dhe 2 të përjashtuara për indeks dimëror bosh.",
            "Grupi 1: 52 kompani; Grupi 2: 27 kompani.",
            "Silhouette final rreth 0.385; PCA 2D ruan rreth 76.98% të variancës.",
        ],
        "StandardScaler dhe RobustScaler krijonin grup singleton për shkak të një vlere ekstreme. Elbow sugjeronte 4 grupe, ndërsa Silhouette 2; u zgjodh k=2 për ndarje më të qartë dhe të balancuar.",
    )

    add_step(
        doc, "8", "Ndërfaqja Streamlit",
        "Bërja e rezultateve të lexueshme dhe të eksplorueshme pa hapur manualisht raportet teknike.",
        [
            "Shtatë pamje: përmbledhje, profil, mot/festa, outlierë, njehsorë/prosumerë, grupe dhe cilësi.",
            "Filtrim sipas kompanisë dhe periudhës; grafikë interaktivë me Plotly.",
            "Lexim Parquet vetëm për kompaninë e zgjedhur dhe cache për performancë.",
        ],
        ["app.py"],
        ["UI-ja u testua me një kompani standarde dhe një prosumer A+/A-."],
        "Ngarkimi i rreth 4.9 milionë rreshtave do ta ngadalësonte aplikacionin. Filtrimi Parquet dhe cache e zgjidhën problemin; u korrigjua edhe prerja vizuale e KPI-ve të mëdha.",
    )

    add_step(
        doc, "9", "Komanda run-all",
        "Ekzekutimi i të gjithë pipeline-it me një komandë dhe pa përzier rezultate të vjetra me një proces të dështuar.",
        [
            "Gjashtë hapat ekzekutohen në .pipeline_staging.",
            "Output-et promovohen vetëm pas suksesit dhe verifikimit të skedarëve të detyrueshëm.",
            "Backup dhe rollback ruajnë rezultatet ekzistuese në rast gabimi.",
            "pipeline_run_summary.json regjistron statusin, kohën, hapat dhe SHA-256 e workbook-ut.",
        ],
        ["python -m enerco_analysis.cli run-all", "outputs/pipeline_run_summary.json"],
        ["Komanda dhe skenari i dështimit u verifikuan me teste automatike."],
        "Rreziku kryesor ishte një gjendje e përzier, ku disa raporte vinin nga inputi i ri dhe disa nga i vjetri. Staging-u transaksional dhe rollback-u e eliminojnë këtë rrezik.",
    )

    doc.add_heading("10. Testimi dhe kontrolli", level=1)
    add_body(doc, "Projekti ka 14 teste automatike që mbulojnë konfigurimin, cilësinë, transformimin, profilet, outlierët, faktorët e jashtëm, grupimin, UI-në dhe run-all.")
    add_code(doc, ".venv\\Scripts\\python.exe -m pytest -q")
    add_body(doc, "Rezultati i verifikuar: 14 teste kalojnë dhe pip check nuk raporton varësi të dëmtuara.")

    doc.add_heading("11. Kufizimet aktuale", level=1)
    add_bullets(doc, [
        "Mungon sektori i biznesit për krahasime zyrtare brenda sektorit.",
        "Mungojnë tensioni, TS Kodi dhe kapaciteti i kontraktuar.",
        "Mungon lokacioni individual; moti i Prishtinës përdoret si proxy.",
        "UI-ja është lokale dhe nuk është publikuar në server.",
        "Parashikimi për konsumatorë të rinj mbetet fazë e ardhshme.",
    ])

    doc.add_heading("12. Rrjedha përfundimtare", level=1)
    add_numbered(doc, [
        "Vendoset workbook-u anonim në data/input.",
        "Kontrollohet konfigurimi dhe ekzekutohet run-all.",
        "Pipeline-i kontrollon cilësinë dhe transformon lexueshëm të dhënat.",
        "Llogariten profilet, outlierët, moti/festat dhe grupet.",
        "Gjenerohen raportet Excel dhe tabelat Parquet.",
        "Rezultatet eksplorohen në Streamlit.",
    ])
    add_code(doc, ".venv\\Scripts\\python.exe -m enerco_analysis.cli run-all\n.venv\\Scripts\\python.exe -m streamlit run app.py")

    doc.add_heading("Përfundim", level=1)
    add_body(
        doc,
        "Projekti është tashmë i përsëritshëm, anonim, i testuar, i dokumentuar dhe i sigurt ndaj ekzekutimeve të dështuara. Ai krijon një bazë të qëndrueshme për rifreskime periodike dhe për fazat e ardhshme, përfshirë pasurimin me metadata dhe parashikimin e konsumit.",
    )

    doc.core_properties.title = "Përmbledhje e projektit të analizës së konsumit - EnerCo"
    doc.core_properties.subject = "Hapat, rezultatet dhe sfidat e projektit Python"
    doc.core_properties.author = "EnerCo"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
